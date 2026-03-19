# -*- coding: utf-8 -*-
"""
DOCX 교수학습-평가 운영 계획 파서

DOCX 파일에서 교수학습-평가 운영 계획 데이터를 추출하여 JSON으로 변환한다.
8열 교수학습-평가 계획 표(시기/시수/누계/단원명/성취기준/평가요소/수업평가방법/비고)를
자동 탐지하여 추출한다.

Usage:
    python parse_docx_input.py --docx input.docx --output data.json
    python parse_docx_input.py --analyze --docx input.docx
"""
import sys
import io
import os
import json
import argparse

# Windows 콘솔 인코딩
if sys.platform == 'win32':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
    except Exception:
        pass

# 교수학습-평가 계획 표의 헤더 키워드 (부분 매칭)
CURRICULUM_HEADER_KEYWORDS = ['시기', '시수', '누계', '단원', '성취기준', '평가', '수업', '비고']
CURRICULUM_HEADER_MIN_MATCH = 5  # 8개 중 5개 이상 매칭되면 해당 표로 인정

# 열 매핑 (헤더 텍스트 → 인덱스)
COL_NAMES = ['period', 'hours', 'cumulative', 'unit', 'standards',
             'eval_elements', 'methods', 'notes']


def parse_docx(docx_path):
    """DOCX 파일에서 교수학습-평가 운영 계획 데이터를 추출."""
    import docx
    from docx.oxml.ns import qn

    doc = docx.Document(docx_path)
    result = {
        'school_info': {
            'subject': '',
            'semester': '',
            'year': '2026',
            'school_name': '',
            'grades': [],
        },
        'curriculum_tables': [],
        'paragraphs': [],
    }

    # 문단에서 기본 정보 추출
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        result['paragraphs'].append(text)

        # 교과명 탐지
        if '교수학습' in text and '평가' in text:
            # "수학과 교수학습 및 평가 운영 계획" 패턴
            for subj in ['수학', '국어', '영어', '사회', '과학', '도덕', '역사',
                         '기술', '가정', '정보', '체육', '음악', '미술', '진로',
                         '한문', '중국어', '일본어', '독일어', '프랑스어', '스페인어']:
                if subj in text:
                    result['school_info']['subject'] = subj
                    break

        # 학기 탐지
        if '1학기' in text:
            result['school_info']['semester'] = '1학기'
        elif '2학기' in text:
            result['school_info']['semester'] = '2학기'

        # 학년도 탐지
        import re
        year_match = re.search(r'(\d{4})학년도', text)
        if year_match:
            result['school_info']['year'] = year_match.group(1)

    # 표에서 데이터 추출
    for table_idx, table in enumerate(doc.tables):
        rows = table.rows
        if len(rows) < 2:
            continue

        # 첫 행(헤더) 텍스트 추출
        header_texts = []
        for cell in rows[0].cells:
            header_texts.append(cell.text.strip())

        # 교수학습-평가 계획 표인지 판별
        match_count = 0
        for keyword in CURRICULUM_HEADER_KEYWORDS:
            for ht in header_texts:
                if keyword in ht:
                    match_count += 1
                    break

        if match_count >= CURRICULUM_HEADER_MIN_MATCH:
            # 교수학습-평가 계획 표 발견
            print(f"  Table {table_idx}: 교수학습-평가 계획 표 발견 ({len(rows)}행)")

            # 열 매핑 결정 (헤더 텍스트 → 실제 열 인덱스)
            col_map = _build_col_map(header_texts)

            # 데이터 행 추출
            data_rows = []
            for row_idx in range(1, len(rows)):
                row = rows[row_idx]
                seen_tcs = set()
                cells_text = []
                for cell in row.cells:
                    tc_id = id(cell._tc)
                    if tc_id in seen_tcs:
                        continue
                    seen_tcs.add(tc_id)
                    cells_text.append(cell.text.strip())

                # 빈 행 건너뛰기
                if all(not ct for ct in cells_text):
                    continue

                # 8열로 정규화
                row_data = _map_row_to_8cols(cells_text, col_map, len(header_texts))
                data_rows.append(row_data)

            if data_rows:
                # 학년 추정: 표 직전 문단에서 "N학년" 검색, 실패 시 순서 기반
                grade = _detect_grade_for_table(table_idx, doc, len(result['curriculum_tables']) + 1)
                result['curriculum_tables'].append({
                    'grade': grade,
                    'rows': data_rows,
                })
                print(f"    {grade}학년: {len(data_rows)}행 추출")

        else:
            # 기본 정보 표 탐지 (학교명, 학급수 등)
            _extract_school_info(table, result['school_info'])

    # paragraphs는 내부용이므로 최종 출력에서 제거
    del result['paragraphs']

    return result


def _build_col_map(header_texts):
    """헤더 텍스트 목록에서 8열 매핑을 구성."""
    col_map = {}
    keywords_priority = [
        ('period', ['시기']),
        ('hours', ['시수']),
        ('cumulative', ['누계']),
        ('unit', ['단원', '단원명']),
        ('standards', ['성취기준', '교육과정']),
        ('eval_elements', ['평가 요소', '평가요소']),
        ('methods', ['수업', '방법', '수업평가', '수업·평가']),
        ('notes', ['비고']),
    ]

    for col_key, keywords in keywords_priority:
        for idx, ht in enumerate(header_texts):
            if idx in col_map.values():
                continue
            for kw in keywords:
                if kw in ht:
                    col_map[col_key] = idx
                    break
            if col_key in col_map:
                break

    return col_map


def _map_row_to_8cols(cells_text, col_map, total_cols):
    """셀 텍스트를 8열 리스트로 매핑."""
    result = [''] * 8
    for i, col_name in enumerate(COL_NAMES):
        if col_name in col_map:
            src_idx = col_map[col_name]
            if src_idx < len(cells_text):
                result[i] = cells_text[src_idx]
        elif i < len(cells_text):
            # 매핑 실패 시 순서대로
            result[i] = cells_text[i] if i < len(cells_text) else ''
    return result


def _detect_grade_for_table(table_idx, doc, fallback_order):
    """표 직전 문단에서 학년을 추정. 실패 시 발견 순서 기반."""
    import re
    from docx.oxml.ns import qn

    # DOCX body에서 이 테이블 직전 문단들을 역방향 탐색
    body = doc.element.body
    current_table_count = 0
    prev_paragraphs = []

    for child in body:
        if child.tag == qn('w:tbl'):
            if current_table_count == table_idx:
                break
            current_table_count += 1
            prev_paragraphs = []  # 새 테이블 구간이면 리셋
        elif child.tag == qn('w:p'):
            text = child.text or ''
            # 모든 텍스트 런도 수집
            for t_elem in child.iter(qn('w:t')):
                if t_elem.text:
                    text += t_elem.text
            prev_paragraphs.append(text.strip())

    # 직전 5개 문단에서 "N학년" 검색 (역순)
    for text in reversed(prev_paragraphs[-5:]):
        m = re.search(r'(\d)학년', text)
        if m:
            return int(m.group(1))

    # 실패 시 순서 기반 (첫 번째 교수학습 표 → 1학년, ...)
    return fallback_order


def _extract_school_info(table, school_info):
    """작은 표에서 학교 기본 정보를 추출."""
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        for i, ct in enumerate(cells):
            if '학교' in ct and i + 1 < len(cells):
                val = cells[i + 1]
                if val and val != ct:
                    school_info['school_name'] = val
            if '교과서' in ct and i + 1 < len(cells):
                val = cells[i + 1]
                if val and val != ct:
                    for g in school_info['grades']:
                        if not g.get('textbook'):
                            g['textbook'] = val
                            break
                    else:
                        school_info['grades'].append({
                            'grade': 1, 'classes': '', 'textbook': val,
                            'hours_per_week': 4,
                        })


def analyze_docx(docx_path):
    """DOCX 구조를 분석하여 요약 출력."""
    import docx
    doc = docx.Document(docx_path)

    print(f"파일: {docx_path}")
    print(f"문단 수: {len(doc.paragraphs)}")
    print(f"표 수: {len(doc.tables)}")
    print()

    for i, table in enumerate(doc.tables):
        rows = table.rows
        if not rows:
            continue
        header = [c.text.strip()[:20] for c in rows[0].cells]
        # 중복 셀 제거
        seen = set()
        unique_header = []
        for h in header:
            if h not in seen:
                unique_header.append(h)
                seen.add(h)
        print(f"Table {i}: {len(rows)}행 x {len(unique_header)}열")
        print(f"  헤더: {unique_header}")
        if len(rows) > 1:
            first_data = [c.text.strip()[:30] for c in rows[1].cells]
            seen2 = set()
            unique_data = []
            for d in first_data:
                if d not in seen2:
                    unique_data.append(d)
                    seen2.add(d)
            print(f"  첫 행: {unique_data}")
        print()


def main():
    parser = argparse.ArgumentParser(description='DOCX 교수학습-평가 운영 계획 파서')
    parser.add_argument('--docx', required=True, help='DOCX 파일 경로')
    parser.add_argument('--output', help='출력 JSON 파일 경로')
    parser.add_argument('--analyze', action='store_true', help='분석 모드 (구조만 출력)')
    args = parser.parse_args()

    if not os.path.exists(args.docx):
        print(f"ERROR: 파일을 찾을 수 없습니다: {args.docx}")
        sys.exit(1)

    if args.analyze:
        analyze_docx(args.docx)
    else:
        data = parse_docx(args.docx)
        output_path = args.output or args.docx.replace('.docx', '_data.json')
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"\n출력: {output_path}")
        print(f"  교과: {data['school_info']['subject']}")
        print(f"  학년: {[ct['grade'] for ct in data['curriculum_tables']]}")
        print(f"  총 행수: {sum(len(ct['rows']) for ct in data['curriculum_tables'])}")


if __name__ == '__main__':
    main()
