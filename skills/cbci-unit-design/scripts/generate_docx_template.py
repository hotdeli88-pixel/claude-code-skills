# -*- coding: utf-8 -*-
"""
전북형 개념기반 탐구학습 단원설계 DOCX 생성 템플릿
────────────────────────────────────────────
이 스크립트는 /cbci-unit-design 스킬에서 자동 생성하는
generate_{unit_name}.py의 기반 템플릿이다.

사용법:
  이 파일을 직접 실행하지 않는다.
  스킬이 단원별 콘텐츠 변수를 채워 새 스크립트를 생성한다.
"""
import sys
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ═══════════════════════════════════════════════
# 헬퍼 함수 (모든 단원 공통)
# ═══════════════════════════════════════════════

def create_doc():
    """새 문서 생성 및 기본 스타일 설정"""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.15
    return doc


def shade(cell, color):
    """셀 배경색 설정"""
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))


def put(cell, text, size=10, bold=False, color=None, align=None, bg=None):
    """셀에 텍스트 입력 (줄바꿈 지원)"""
    if bg:
        shade(cell, bg)
    for p in list(cell.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for ch in list(first._element):
        t = ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag
        if t != 'pPr':
            first._element.remove(ch)
    for i, line in enumerate(text.split('\n')):
        p = first if i == 0 else cell.add_paragraph()
        if align:
            p.alignment = align
        r = p.add_run(line)
        r.font.name = '맑은 고딕'
        r.font.size = Pt(size)
        r.font.bold = bold
        if color:
            r.font.color.rgb = (RGBColor.from_string(color)
                                if isinstance(color, str) else color)


def hdr(cell, text, size=10):
    """파란 헤더 셀 (흰 글자)"""
    put(cell, text, size=size, bold=True, color='FFFFFF', bg='2E75B6',
        align=WD_ALIGN_PARAGRAPH.CENTER)


def sub(cell, text, size=10):
    """연파랑 서브헤더 셀"""
    put(cell, text, size=size, bold=True, bg='D6E4F0',
        align=WD_ALIGN_PARAGRAPH.CENTER)


def label(cell, text, size=10):
    """라벨 셀 (연파랑 배경)"""
    put(cell, text, size=size, bold=True, bg='D6E4F0')


def set_col_widths(table, widths_cm):
    """테이블 열 너비 설정"""
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def title(doc, text, level=0):
    """제목 추가"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = '맑은 고딕'
    if level == 0:
        r.font.size = Pt(16)
        r.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    elif level == 2:
        r.font.size = Pt(11)
        r.font.bold = True
    return p


def spacer(doc):
    """빈 줄"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════
# 테이블 생성 함수 (모든 단원 공통)
# ═══════════════════════════════════════════════

def create_overview_table(doc, data_list):
    """Table 1: 유형 개요 (key-value 쌍 리스트)"""
    t = doc.add_table(rows=len(data_list), cols=2)
    t.style = 'Table Grid'
    set_col_widths(t, [4, 13])
    for i, (k, v) in enumerate(data_list):
        label(t.rows[i].cells[0], k)
        put(t.rows[i].cells[1], v)
    return t


def create_curriculum_table(doc, data_list):
    """Table 2: 교육과정 분석 (key-value, 특수 서식 포함)"""
    t = doc.add_table(rows=len(data_list), cols=2)
    t.style = 'Table Grid'
    set_col_widths(t, [4, 13])
    for i, (k, v, *opts) in enumerate(data_list):
        label(t.rows[i].cells[0], k)
        kwargs = {}
        if opts:
            kwargs = opts[0]
        put(t.rows[i].cells[1], v, **kwargs)
    return t


def create_subtopic_table(doc, subtopics):
    """Table 3: 소주제 구성
    subtopics: [(이름, 차시, 모형, 일반화), ...]
    """
    t = doc.add_table(rows=len(subtopics)+1, cols=4)
    t.style = 'Table Grid'
    set_col_widths(t, [4, 2, 3, 8])
    for j, h in enumerate(['소주제', '차시', '학습모형', '소주제 일반화']):
        hdr(t.rows[0].cells[j], h)
    for i, (name, hours, model, gen) in enumerate(subtopics):
        put(t.rows[i+1].cells[0], name, bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER)
        put(t.rows[i+1].cells[1], hours, align=WD_ALIGN_PARAGRAPH.CENTER)
        put(t.rows[i+1].cells[2], model, align=WD_ALIGN_PARAGRAPH.CENTER)
        put(t.rows[i+1].cells[3], gen)
    return t


def create_inquiry_table(doc, rows_data):
    """Table 4: 탐구 질문 종합
    rows_data: [(구분, 유형, 질문, {옵션}), ...]
    """
    t = doc.add_table(rows=len(rows_data)+1, cols=3)
    t.style = 'Table Grid'
    set_col_widths(t, [3, 3, 11])
    hdr(t.rows[0].cells[0], '구분')
    hdr(t.rows[0].cells[1], '유형')
    hdr(t.rows[0].cells[2], '질문')
    for i, (sec, qtype, question, *opts) in enumerate(rows_data):
        sub(t.rows[i+1].cells[0], sec)
        sub(t.rows[i+1].cells[1], qtype)
        kwargs = opts[0] if opts else {}
        put(t.rows[i+1].cells[2], question, **kwargs)
    return t


def create_eval_table(doc, eval_data):
    """Table 5: 평가 계획
    eval_data: [(차원, 차시, 내용, 방법), ...]
    """
    t = doc.add_table(rows=len(eval_data)+1, cols=4)
    t.style = 'Table Grid'
    set_col_widths(t, [3, 2, 8, 4])
    for j, h in enumerate(['차원', '차시', '평가 내용', '방법']):
        hdr(t.rows[0].cells[j], h)
    for i, (dim, hour, content, method) in enumerate(eval_data):
        sub(t.rows[i+1].cells[0], dim)
        put(t.rows[i+1].cells[1], hour, align=WD_ALIGN_PARAGRAPH.CENTER)
        put(t.rows[i+1].cells[2], content)
        put(t.rows[i+1].cells[3], method, align=WD_ALIGN_PARAGRAPH.CENTER)
    return t


def create_subtopic_header_table(doc, gen, key_q, inquiry_q, concept_focus):
    """소주제 헤더 테이블"""
    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    set_col_widths(t, [4, 13])
    label(t.rows[0].cells[0], '소주제 일반화')
    put(t.rows[0].cells[1], gen, bold=True)
    label(t.rows[1].cells[0], '핵심 질문')
    put(t.rows[1].cells[1], key_q)
    label(t.rows[2].cells[0], '탐구 질문\n(사/개/논)')
    put(t.rows[2].cells[1], inquiry_q)
    label(t.rows[3].cells[0], '핵심개념 초점')
    put(t.rows[3].cells[1], concept_focus)
    return t


def create_activity_table(doc, lessons, model_type='discovery'):
    """소주제 활동 테이블
    lessons: [(차시, 탐구단계, 활동내용), ...]
    model_type: 'discovery'(발견), 'principle'(원리), 'process'(과정)
    """
    t = doc.add_table(rows=len(lessons)+1, cols=3)
    t.style = 'Table Grid'
    set_col_widths(t, [2, 3, 12])
    hdr(t.rows[0].cells[0], '차시')
    hdr(t.rows[0].cells[1], '탐구 단계')
    hdr(t.rows[0].cells[2], '활동 내용')
    for i, (hour, stage, content) in enumerate(lessons):
        sub(t.rows[i+1].cells[0], hour)
        sub(t.rows[i+1].cells[1], stage)
        put(t.rows[i+1].cells[2], content)
    return t


# ═══════════════════════════════════════════════
# 전체 문서 조립 함수
# ═══════════════════════════════════════════════

def build_unit_design(
    unit_name,
    overview_data,
    curriculum_data,
    subtopics,
    inquiry_rows,
    eval_data,
    subtopic_details,
    output_file=None
):
    """
    전체 단원설계 DOCX를 생성한다.

    Parameters
    ----------
    unit_name : str
        단원명 (예: "유리수와 순환소수")
    overview_data : list of (key, value)
        유형 개요 데이터
    curriculum_data : list of (key, value, {opts})
        교육과정 분석 데이터
    subtopics : list of (name, hours, model, generalization)
        소주제 구성 데이터
    inquiry_rows : list of (section, type, question, {opts})
        탐구 질문 종합 데이터
    eval_data : list of (dimension, hour, content, method)
        평가 계획 데이터
    subtopic_details : list of dict
        각 소주제의 상세 정보:
        {
            'title': '소주제 N 「제목」 모형명 설계 (X~Y차시)',
            'gen': 일반화,
            'key_q': 핵심질문,
            'inquiry_q': 탐구질문,
            'concept_focus': 핵심개념 초점,
            'model_type': 'discovery'|'principle'|'process',
            'lessons': [(차시, 탐구단계, 활동내용), ...]
        }
    output_file : str, optional
        출력 파일명. None이면 단원명 기반 자동 생성
    """
    if output_file is None:
        safe_name = unit_name.replace(' ', '').replace('/', '_')
        output_file = f'{safe_name}_단원설계.docx'

    doc = create_doc()

    # 문서 제목
    title(doc, '전북형 개념기반 탐구학습 단원설계', level=0)
    p = doc.add_paragraph()
    r = p.add_run(unit_name)
    r.font.name = '맑은 고딕'
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer(doc)

    # Table 1: 유형 개요
    title(doc, '1. 유형 개요', level=1)
    create_overview_table(doc, overview_data)
    print("[OK] Table 1: 유형 개요")
    spacer(doc)

    # Table 2: 교육과정 분석
    title(doc, '2. 교육과정 분석', level=1)
    create_curriculum_table(doc, curriculum_data)
    print("[OK] Table 2: 교육과정 분석")
    spacer(doc)

    # Table 3: 소주제 구성
    title(doc, '3. 소주제 구성', level=1)
    create_subtopic_table(doc, subtopics)
    print("[OK] Table 3: 소주제 구성")
    spacer(doc)

    # Table 4: 탐구 질문 종합
    title(doc, '4. 탐구 질문 종합', level=1)
    create_inquiry_table(doc, inquiry_rows)
    print("[OK] Table 4: 탐구 질문 종합")
    spacer(doc)

    # Table 5: 평가 계획
    title(doc, '5. 평가 계획 (3차원)', level=1)
    create_eval_table(doc, eval_data)
    print("[OK] Table 5: 평가 계획")
    spacer(doc)

    # 소주제별 상세 (헤더 + 활동)
    for idx, st in enumerate(subtopic_details):
        title(doc, st['title'], level=1)
        create_subtopic_header_table(
            doc, st['gen'], st['key_q'], st['inquiry_q'], st['concept_focus'])
        spacer(doc)
        create_activity_table(doc, st['lessons'], st.get('model_type', 'discovery'))
        print(f"[OK] 소주제 {idx+1} 상세")
        spacer(doc)

    # 저장
    doc.save(output_file)
    print(f"\n[완료] 최종 파일 저장: {output_file}")
    return output_file
